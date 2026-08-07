import os
import logging
import docx
from fpdf import FPDF
from src.tools.file_ops.permissions import FilePermissions
from typing import Callable

log = logging.getLogger(__name__)

class FileWriters:
    """Handles creating and writing files in various formats."""

    def __init__(self, ask_callback: Callable[[str], bool]):
        # The callback used to ask the user for ALWAYS_ASK permissions
        self.ask_callback = ask_callback

    def write_file(self, file_path: str, content: str) -> str:
        """Writes content to a file, creating it if necessary."""
        # 1. Validation for Overwrites vs Creates
        operation = "overwrite" if os.path.exists(file_path) else "create"
        
        if not FilePermissions.validate(operation, file_path, self.ask_callback):
            return f"Security Error: Permission denied to {operation} {file_path}."

        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            # Ensure parent directories exist
            os.makedirs(os.path.dirname(os.path.abspath(file_path)) or '.', exist_ok=True)
            
            if ext in ['.txt', '.md', '.csv', '.json', '.py']:
                return self._write_text(file_path, content)
            elif ext == '.pdf':
                return self._write_pdf(file_path, content)
            elif ext == '.docx':
                return self._write_docx(file_path, content)
            else:
                return f"Error: Unsupported file format {ext} for writing."
        except Exception as e:
            log.error(f"Failed to write {file_path}: {e}")
            return f"Error writing file: {e}"

    def _write_text(self, file_path: str, content: str) -> str:
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return f"Successfully saved text to {file_path}."

    def _write_pdf(self, file_path: str, content: str) -> str:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        # Handle encoding for FPDF
        content = content.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 10, txt=content)
        pdf.output(file_path)
        return f"Successfully generated PDF document at {file_path}."

    def _write_docx(self, file_path: str, content: str) -> str:
        doc = docx.Document()
        
        # Simple parsing for headings based on Markdown
        for line in content.split('\n'):
            if line.startswith('# '):
                doc.add_heading(line.replace('# ', ''), 0)
            elif line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), 1)
            elif line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), 2)
            else:
                doc.add_paragraph(line)
                
        doc.save(file_path)
        return f"Successfully generated Word document at {file_path}."
